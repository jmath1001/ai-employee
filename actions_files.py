import os
from docx import Document

WORKSPACE = r"C:\Users\david\my-ai-agent\bot_files"
if not os.path.exists(WORKSPACE):
    os.makedirs(WORKSPACE)


def write_text(filename, content):
    path = os.path.join(WORKSPACE, filename)
    with open(path, 'w', encoding='utf-8') as f:
        f.write(content)
    return f"Successfully wrote to {filename}"


def read_any_file(filename):
    path = os.path.join(WORKSPACE, filename)
    if not os.path.exists(path):
        return f"File not found: {filename}"

    if filename.endswith('.docx'):
        doc = Document(path)
        return "\n".join([p.text for p in doc.paragraphs])
    else:
        with open(path, 'r', encoding='utf-8') as f:
            return f.read()


def manage_word_doc(filename, title=None, content=None, mode="create"):
    path = os.path.join(WORKSPACE, filename)
    if not filename.endswith('.docx'):
        path += '.docx'

    if mode == "create":
        doc = Document()
        if title:
            doc.add_heading(title, 0)

        # Parse content for markdown-style formatting
        for line in content.split('\n'):
            line = line.strip()
            if not line:
                continue
            elif line.startswith('### '):
                doc.add_heading(line[4:], level=3)
            elif line.startswith('## '):
                doc.add_heading(line[3:], level=2)
            elif line.startswith('# '):
                doc.add_heading(line[2:], level=1)
            elif line.startswith('- ') or line.startswith('* '):
                doc.add_paragraph(line[2:], style='List Bullet')
            elif line[0].isdigit() and '. ' in line:
                doc.add_paragraph(line, style='List Number')
            else:
                doc.add_paragraph(line)

        doc.save(path)
        return f"Word doc {filename} created."

    elif mode == "append":
        doc = Document(path)
        for line in content.split('\n'):
            line = line.strip()
            if not line:
                continue
            elif line.startswith('### '):
                doc.add_heading(line[4:], level=3)
            elif line.startswith('## '):
                doc.add_heading(line[3:], level=2)
            elif line.startswith('# '):
                doc.add_heading(line[2:], level=1)
            elif line.startswith('- ') or line.startswith('* '):
                doc.add_paragraph(line[2:], style='List Bullet')
            elif line[0].isdigit() and '. ' in line:
                doc.add_paragraph(line, style='List Number')
            else:
                doc.add_paragraph(line)
        doc.save(path)
        return f"Added content to {filename}."

# ADD THIS PART - IT WAS MISSING!
def list_files():
    files = sorted(os.listdir(WORKSPACE))
    return f"Files: {', '.join(files)}" if files else "The workspace is empty."