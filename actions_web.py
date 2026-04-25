import os
import re
import time
import urllib.parse
import urllib.request
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from playwright.sync_api import sync_playwright
from duckduckgo_search import DDGS

WORKSPACE = r"C:\Users\david\my-ai-agent\bot_files"

# ──────────────────────────────────────────────
# HELPERS
# ──────────────────────────────────────────────

def _make_browser(p):
    """Launch a stealthy Chromium instance."""
    browser = p.chromium.launch(
        headless=True,
        args=["--disable-blink-features=AutomationControlled"]
    )
    context = browser.new_context(
        user_agent=(
            "Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
            "AppleWebKit/537.36 (KHTML, like Gecko) "
            "Chrome/122.0.0.0 Safari/537.36"
        )
    )
    context.add_init_script(
        "Object.defineProperty(navigator, 'webdriver', { get: () => undefined });"
    )
    return browser, context


def _safe_goto(page, url, timeout=15000):
    """Navigate with fallback."""
    try:
        page.goto(url, wait_until="domcontentloaded", timeout=timeout)
        time.sleep(1)
        return True
    except Exception as e:
        print(f"[WARN] Could not load {url}: {e}")
        return False


def _clean_page_text(page, max_chars=4000):
    """Strip nav/footer junk and return body text."""
    try:
        page.evaluate(
            "document.querySelectorAll("
            "'nav, footer, .ad, .popup, .modal, .cookie-banner,"
            " .sidebar, .related-posts, script, style'"
            ").forEach(el => el.remove())"
        )
        text = page.locator("body").inner_text()
        return text[:max_chars].strip()
    except:
        return ""


def _save_docx(filename: str, title: str, sections: list[dict]) -> str:
    """
    Save a structured Word document.

    sections = [
        {"heading": "Section Title", "body": "paragraph text"},
        {"heading": "Sources", "bullets": ["url1", "url2"]},
        ...
    ]
    Returns the full file path.
    """
    os.makedirs(WORKSPACE, exist_ok=True)
    safe_name = re.sub(r"[^\w]", "_", filename)[:50]
    path = os.path.join(WORKSPACE, f"{safe_name}.docx")

    doc = Document()

    # Title
    title_para = doc.add_heading(title, level=0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for section in sections:
        heading = section.get("heading")
        body = section.get("body")
        bullets = section.get("bullets", [])

        if heading:
            doc.add_heading(heading, level=1)

        if body:
            doc.add_paragraph(body)

        for bullet in bullets:
            doc.add_paragraph(bullet, style="List Bullet")

    doc.save(path)
    print(f"[OK] Saved: {path}")
    return path


def _ai_summarize(client, prompt: str, context: str, max_ctx=10000) -> str:
    """Call local Qwen2.5 to summarize. Return fallback if AI fails."""
    if not context or len(context.strip()) < 20:
        return "[No content extracted to summarize. Check if websites are accessible.]"
    
    try:
        full_prompt = f"{prompt}\n\nDATA:\n{context[:max_ctx]}"
        response = client.chat.completions.create(
            model="qwen2.5",
            messages=[{"role": "user", "content": full_prompt}],
            extra_body={"options": {"num_ctx": 12000}}
        )
        return response.choices[0].message.content.strip()
    except Exception as e:
        print(f"[WARN] AI summarize failed: {e}")
        return f"[Summarization unavailable: {e}. Raw content excerpt:] {context[:500]}"


def _fallback_search_links(query: str, max_results=5) -> list[str]:
    """Fallback search strategy when browser selectors fail."""
    links = []

    # 1) DDGS API attempts (retry + backend variations)
    for attempt in range(3):
        for backend in ("auto", "lite", "html"):
            try:
                ddgs_results = list(DDGS().text(query, max_results=max_results, backend=backend))
                links = [r.get("href") for r in ddgs_results if r.get("href")]
                if links:
                    return links[:max_results]
            except Exception as e:
                print(f"   DDGS fallback failed (attempt {attempt + 1}, backend {backend}): {e}")
        time.sleep(0.5)

    if links:
        return links[:max_results]

    # 2) Bing HTML fallback (no JS selector dependency)
    try:
        encoded = urllib.parse.quote_plus(query)
        url = f"https://www.bing.com/search?q={encoded}&count={max_results}"
        req = urllib.request.Request(
            url,
            headers={
                "User-Agent": (
                    "Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
                    "AppleWebKit/537.36 (KHTML, like Gecko) "
                    "Chrome/122.0.0.0 Safari/537.36"
                )
            },
        )
        with urllib.request.urlopen(req, timeout=12) as resp:
            html = resp.read().decode("utf-8", errors="ignore")

        candidates = re.findall(r'<a href="(https?://[^"]+)"', html)
        blocked = ("bing.com", "microsoft.com", "go.microsoft.com")
        cleaned = []
        for href in candidates:
            if any(b in href for b in blocked):
                continue
            if href not in cleaned:
                cleaned.append(href)
            if len(cleaned) >= max_results:
                break
        links = cleaned
    except Exception as e:
        print(f"   Bing fallback failed: {e}")

    return links[:max_results]


# ──────────────────────────────────────────────
# 1. DEEP RESEARCH  (Edge/Selenium browser)
# ──────────────────────────────────────────────

def deep_research(query: str, client) -> str:
    """
    Use Edge (Selenium) to search Bing, visit top result pages, then write a Word doc.
    Falls back to DDGS API if browser link extraction fails.
    Always returns a Word doc even on partial failure.
    """
    from selenium.webdriver.common.by import By
    from selenium.webdriver.support.ui import WebDriverWait
    from selenium.webdriver.support import expected_conditions as EC
    from selenium.common.exceptions import TimeoutException, WebDriverException
    from browser_engine import EdgeEngine

    driver = EdgeEngine.get_driver()
    findings = []
    links = []

    try:
        print(f"[SEARCH] {query}")
        search_url = f"https://www.bing.com/search?q={urllib.parse.quote_plus(query)}"
        driver.get(search_url)
        time.sleep(2)

        # Extract result links from Bing
        try:
            result_elements = driver.find_elements(By.CSS_SELECTOR, "li.b_algo h2 a")
            links = [el.get_attribute("href") for el in result_elements[:6] if el.get_attribute("href")]
        except Exception as e:
            print(f"   Bing selector failed: {e}")

        if not links:
            print("   Bing selectors returned nothing, trying DDGS fallback...")
            links = _fallback_search_links(query, max_results=5)

        print(f"   Found {len(links)} links to scrape")

        for link in links:
            if not link:
                continue
            # Keep Bing redirect URLs (bing.com/ck/...) because they resolve to real targets.
            if (
                "bing.com" in link
                and "bing.com/ck/" not in link
                and "bing.com/aclick" not in link
            ):
                continue
            if any(s in link for s in ("google.com/search", "youtube.com")):
                continue

            try:
                print(f"[SCRAPE] {link}")
                driver.get(link)
                WebDriverWait(driver, 10).until(
                    EC.presence_of_element_located((By.TAG_NAME, "body"))
                )
                time.sleep(1.5)
                final_url = driver.current_url

                # If we are still on a search/redirect domain, skip this result.
                if any(
                    x in final_url
                    for x in ("bing.com/search", "google.com/search", "duckduckgo.com")
                ):
                    print(f"   Skipping non-article destination: {final_url}")
                    continue

                # Remove nav/footer noise via JS
                driver.execute_script(
                    "document.querySelectorAll('nav,footer,.ad,.popup,.modal,.cookie-banner,.sidebar,script,style')"
                    ".forEach(el => el.remove());"
                )
                body = driver.find_element(By.TAG_NAME, "body")
                text = body.text[:3000].strip()

                if text and len(text) > 80:
                    findings.append({"url": final_url or link, "content": text})

                if len(findings) >= 3:
                    break

            except (TimeoutException, WebDriverException) as e:
                print(f"   Failed to load {link}: {e}")
                continue

    except Exception as e:
        print(f"[WARN] Search phase error: {e}")

    # Build Word doc from findings
    if findings:
        context_block = ""
        for i, f in enumerate(findings):
            context_block += f"\n--- SOURCE #{i+1}: {f['url']} ---\n{f['content']}"

        summary = _ai_summarize(
            client,
            prompt=(
                f"You are a research analyst. Synthesize everything known about '{query}' "
                f"from the sources below. Include: key facts, dates, prices, names, "
                f"controversies, and any data worth noting. Use markdown tables where helpful."
            ),
            context=context_block
        )

        sections = [
            {"heading": "Summary", "body": summary},
            {"heading": "Sources", "bullets": [f["url"] for f in findings]},
        ]
    else:
        sections = [
            {"heading": "Search Failed", "body": (
                f"Could not extract content for '{query}'. "
                "Bing and fallback search backends returned no usable results. "
                "Websites may be blocking automated access. Try again or search manually."
            )}
        ]

    path = _save_docx(
        filename=f"research_{query.replace(' ', '_')}",
        title=f"Research Report: {query}",
        sections=sections
    )

    return f"FILE_SIGNAL|{path}"


# ──────────────────────────────────────────────
# 2. REDDIT DIVER
# ──────────────────────────────────────────────

def reddit_dive(query: str, client, max_posts=5) -> str:
    """
    Search Reddit for a topic, open threads, scrape top comments,
    and produce a Word doc with real opinions.
    """
    with sync_playwright() as p:
        browser, context = _make_browser(p)
        page = context.new_page()

        try:
            search_url = f"https://www.reddit.com/search/?q={query.replace(' ', '+')}&sort=relevance&t=year"
            print(f"[REDDIT] {query}")
            if not _safe_goto(page, search_url):
                browser.close()
                return "Could not reach Reddit."

            time.sleep(2)

            # Grab post links
            post_links = page.evaluate("""
                () => [...document.querySelectorAll('a[data-testid="post-title"]')]
                    .map(a => 'https://www.reddit.com' + a.getAttribute('href'))
                    .filter(h => h.includes('/comments/'))
                    .slice(0, 8)
            """)

            if not post_links:
                # fallback selector for new Reddit layout
                post_links = page.evaluate("""
                    () => [...document.querySelectorAll('a[slot="full-post-link"]')]
                        .map(a => 'https://www.reddit.com' + a.getAttribute('href'))
                        .slice(0, 8)
                """)

            print(f"   Found {len(post_links)} posts")

            threads = []

            for post_url in post_links[:max_posts]:
                print(f"   Reading: {post_url}")
                if not _safe_goto(page, post_url):
                    continue

                time.sleep(1.5)

                # Post title
                try:
                    post_title = page.locator("h1").first.inner_text()
                except:
                    post_title = post_url

                # OP body
                try:
                    op_text = page.locator('[data-testid="post-content"]').inner_text()[:800]
                except:
                    op_text = ""

                # Top comments (grab first 10 visible)
                try:
                    comments = page.evaluate("""
                        () => [...document.querySelectorAll('[data-testid="comment"]')]
                            .slice(0, 10)
                            .map(c => c.innerText.trim().slice(0, 300))
                            .filter(t => t.length > 20)
                    """)
                except:
                    comments = []

                if not comments:
                    # fallback for new layout
                    try:
                        comments = page.evaluate("""
                            () => [...document.querySelectorAll('shreddit-comment')]
                                .slice(0, 10)
                                .map(c => c.innerText.trim().slice(0, 300))
                                .filter(t => t.length > 20)
                        """)
                    except:
                        comments = []

                threads.append({
                    "title": post_title,
                    "url": post_url,
                    "op": op_text,
                    "comments": comments
                })

            # Build context block (even if minimal)
            if threads:
                context_block = ""
                for t in threads:
                    context_block += f"\n\n=== POST: {t['title']} ===\n{t['op']}\n"
                    if t.get('comments'):
                        context_block += "TOP COMMENTS:\n" + "\n---\n".join(t["comments"])
            else:
                context_block = f"[Reddit search for '{query}' returned no extractable threads. Site may be blocking automated access.]"

            summary = _ai_summarize(
                client,
                prompt=(
                    f"Analyze Reddit's community opinions on '{query}'. "
                    f"Summarize: the main consensus, key debates, top concerns, "
                    f"notable personal experiences, and any standout advice or warnings. "
                    f"Write in a clear, journalistic style."
                ),
                context=context_block
            )

            sections = [
                {"heading": "Reddit Community Analysis", "body": summary},
            ]
            if threads:
                sections.append({"heading": "Threads Reviewed"})
                for t in threads:
                    sections.append({"body": f"• {t['title']}\n  {t['url']}"})
            else:
                sections.append({"body": "No threads extracted. Reddit may be blocking the search."})

            path = _save_docx(
                filename=f"reddit_{query.replace(' ', '_')}",
                title=f"Reddit Deep Dive: {query}",
                sections=sections
            )

            browser.close()
            return f"FILE_SIGNAL|{path}"

        except Exception as e:
            browser.close()
            print(f"[WARN] Reddit Error: {e}")
            path = _save_docx(
                filename=f"reddit_{query.replace(' ', '_')}",
                title=f"Reddit Deep Dive: {query}",
                sections=[{"heading": "Error", "body": f"Reddit dive failed: {e}"}]
            )
            return f"FILE_SIGNAL|{path}"


# ──────────────────────────────────────────────
# 3. NEWS AGGREGATOR  (multi-source)
# ──────────────────────────────────────────────

NEWS_SOURCES = [
    ("BBC",         "https://www.bbc.com/search?q={q}"),
    ("Reuters",     "https://www.reuters.com/search/news?blob={q}"),
    ("AP News",     "https://apnews.com/search?q={q}"),
    ("The Guardian","https://www.theguardian.com/search?q={q}"),
    ("Al Jazeera",  "https://www.aljazeera.com/search/{q}"),
]

def aggregate_news(query: str, client, num_sources=3) -> str:
    """
    Hits multiple news sources, grabs article text, and synthesizes
    a balanced multi-perspective Word doc.
    """
    with sync_playwright() as p:
        browser, context = _make_browser(p)
        page = context.new_page()

        findings = []

        for source_name, url_template in NEWS_SOURCES[:num_sources]:
            url = url_template.replace("{q}", query.replace(" ", "+"))
            print(f"[NEWS] {source_name}: {url}")

            if not _safe_goto(page, url):
                continue

            time.sleep(1.5)

            # Try to find first article link on the search results page
            try:
                article_link = page.evaluate("""
                    () => {
                        const a = document.querySelector('article a, h3 a, h2 a, .headline a');
                        return a ? a.href : null;
                    }
                """)
            except:
                article_link = None

            # If we found an article, visit it
            text = ""
            article_url = url
            if article_link and article_link.startswith("http"):
                print(f"   Article: {article_link}")
                if _safe_goto(page, article_link):
                    article_url = article_link
                    text = _clean_page_text(page, max_chars=3500)

            if not text:
                # Fall back to the search results page itself
                text = _clean_page_text(page, max_chars=3500)

            if text:
                findings.append({
                    "source": source_name,
                    "url": article_url,
                    "content": text
                })

        browser.close()

        if not findings:
            # Return error doc instead of plain text
            path = _save_docx(
                filename=f"news_{query.replace(' ', '_')}",
                title=f"News Report: {query}",
                sections=[{"heading": "No Results", "body": f"Could not find news coverage for '{query}'. News sites may be blocking automated access."}]
            )
            return f"FILE_SIGNAL|{path}"

        context_block = ""
        for f in findings:
            context_block += f"\n\n=== {f['source'].upper()} ({f['url']}) ===\n{f['content']}"

        summary = _ai_summarize(
            client,
            prompt=(
                f"You are a senior news editor. Synthesize coverage of '{query}' from multiple sources. "
                f"Write: (1) A factual summary of what happened, (2) Key quotes/claims from each source, "
                f"(3) Any differences in how sources frame the story, (4) Important missing context. "
                f"Be objective and note any bias."
            ),
            context=context_block
        )

        sections = [
            {"heading": "News Summary", "body": summary},
            {"heading": "Sources Reviewed"},
        ]
        for f in findings:
            sections.append({"body": f"• {f['source']}: {f['url']}"})

        path = _save_docx(
            filename=f"news_{query}",
            title=f"News Report: {query}",
            sections=sections
        )

        return f"FILE_SIGNAL|{path}"


# ──────────────────────────────────────────────
# 4. TWITTER / X SCRAPER  (Nitter fallback)
# ──────────────────────────────────────────────

NITTER_INSTANCES = [
    "https://nitter.privacydev.net",
    "https://nitter.poast.org",
    "https://nitter.1d4.us",
]

def scrape_twitter(query: str, client, max_tweets=30) -> str:
    """
    Scrapes tweets via Nitter (public Twitter mirror — no login needed).
    Summarizes sentiment and key opinions into a Word doc.
    """
    with sync_playwright() as p:
        browser, context = _make_browser(p)
        page = context.new_page()

        tweets_collected = []
        working_instance = None

        for instance in NITTER_INSTANCES:
            search_url = f"{instance}/search?f=tweets&q={query.replace(' ', '+')}"
            print(f"[NITTER] Trying: {instance}")
            try:
                page.goto(search_url, wait_until="domcontentloaded", timeout=12000)
                time.sleep(2)

                # Check if page loaded properly
                if "timeline-item" in page.content() or "tweet-content" in page.content():
                    working_instance = instance
                    print(f"   [OK] {instance} is working")
                    break
            except:
                print(f"   [DOWN] {instance}")
                continue

        if not working_instance:
            browser.close()
            return f"All Nitter instances down. Try again later."

        # Scrape tweets
        try:
            raw_tweets = page.evaluate("""
                () => [...document.querySelectorAll('.tweet-content')]
                    .map(t => t.innerText.trim())
                    .filter(t => t.length > 10)
            """)

            # Also grab usernames + stats for context
            tweet_meta = page.evaluate("""
                () => [...document.querySelectorAll('.timeline-item')]
                    .slice(0, 40)
                    .map(item => {
                        const user = item.querySelector('.username')?.innerText || '';
                        const text = item.querySelector('.tweet-content')?.innerText || '';
                        const stats = item.querySelector('.tweet-stats')?.innerText || '';
                        return user + ': ' + text + ' [' + stats.replace(/\\n/g, ' ') + ']';
                    })
                    .filter(t => t.length > 15)
            """)
        except:
            raw_tweets = []
            tweet_meta = []

        # Scroll and grab more
        for _ in range(2):
            try:
                page.evaluate("window.scrollTo(0, document.body.scrollHeight)")
                time.sleep(1.5)
                more = page.evaluate("""
                    () => [...document.querySelectorAll('.tweet-content')]
                        .map(t => t.innerText.trim())
                        .filter(t => t.length > 10)
                """)
                raw_tweets = list(dict.fromkeys(raw_tweets + more))  # dedupe
                if len(raw_tweets) >= max_tweets:
                    break
            except:
                break

        browser.close()

        tweets_collected = (tweet_meta or raw_tweets)[:max_tweets]

        if not tweets_collected:
            path = _save_docx(
                filename=f"twitter_{query.replace(' ', '_')}",
                title=f"Twitter Analysis: {query}",
                sections=[{"heading": "No Results", "body": f"No tweets found for '{query}'. Nitter instances may be down or Twitter has blocked access."}]
            )
            return f"FILE_SIGNAL|{path}"

        print(f"   [INFO] Collected {len(tweets_collected)} tweets")

        tweet_block = "\n---\n".join(tweets_collected)

        summary = _ai_summarize(
            client,
            prompt=(
                f"Analyze Twitter/X discussion about '{query}'. "
                f"Summarize: (1) Overall sentiment (positive/negative/mixed), "
                f"(2) The main opinions and arguments people are making, "
                f"(3) Any notable accounts or viral takes, "
                f"(4) Trending sub-topics or hashtags mentioned. "
                f"Be concise and quote interesting tweets directly."
            ),
            context=tweet_block
        )

        sections = [
            {"heading": "Twitter Sentiment Summary", "body": summary},
            {"heading": f"Sample Tweets ({len(tweets_collected)} collected)"},
        ]
        # Add a sample of raw tweets
        for t in tweets_collected[:10]:
            sections.append({"body": t[:200]})

        path = _save_docx(
            filename=f"twitter_{query.replace(' ', '_')}",
            title=f"Twitter Analysis: {query}",
            sections=sections
        )

        return f"FILE_SIGNAL|{path}"